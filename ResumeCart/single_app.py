from dotenv import load_dotenv
import base64
import streamlit as st
import os
import io
from PIL import Image
import pdf2image
import google.generativeai as genai
import matplotlib.pyplot as plt
import mplcursors
import re
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
import google.api_core.exceptions

# Load environment variables
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# ---- FUNCTIONS ----

def get_gemini_response(input_prompt, pdf_content, jd, retries=3):
    model = genai.GenerativeModel('gemini-2.0-flash')
    
    for attempt in range(retries):
        try:
            response = model.generate_content([input_prompt, pdf_content[0], jd])
            return response.text
        except google.api_core.exceptions.ResourceExhausted as e:
            retry_seconds = 60  # Default
            match = re.search(r'retry_delay\s*{\s*seconds:\s*(\d+)', str(e))
            if match:
                retry_seconds = int(match.group(1))
            st.warning(f"⚠️ Quota limit reached (Attempt {attempt + 1}/{retries}). Retrying in {retry_seconds} seconds...")
            if attempt < retries - 1:
                with st.spinner(f"⏳ Waiting for {retry_seconds} seconds..."):
                    time.sleep(retry_seconds)
            else:
                st.error("❌ Failed after maximum retry attempts due to quota exhaustion.")
                raise e

def input_pdf_setup(uploaded_file):
    if uploaded_file is not None:
        images = pdf2image.convert_from_bytes(uploaded_file.read())
        first_page = images[0]
        img_byte_arr = io.BytesIO()
        first_page.save(img_byte_arr, format='JPEG')
        img_byte_arr = img_byte_arr.getvalue()

        pdf_parts = [{
            "mime_type": "image/jpeg",
            "data": base64.b64encode(img_byte_arr).decode()
        }]
        return pdf_parts
    else:
        raise FileNotFoundError("No file uploaded")

def extract_percentage(text):
    match = re.search(r'(\d{1,3})\s*%', text)
    return int(match.group(1)) if match else 0

def extract_missing_keywords(text):
    match = re.search(r'keywords missing[:\-]?\s*(.*?)(?:Final thoughts|$)', text, re.IGNORECASE | re.DOTALL)
    return match.group(1).strip() if match else "N/A"

def show_percentage_chart(percentage):
    match_data = [percentage, 100 - percentage]
    match_labels = [f"Match: {percentage}%", "Remaining"]
    match_colors = ['#4CAF50', '#FFCDD2']

    fig, ax = plt.subplots(figsize=(5, 5))
    wedges, texts, autotexts = ax.pie(
        match_data,
        labels=match_labels,
        colors=match_colors,
        autopct='%1.1f%%',
        startangle=90,
        wedgeprops={'edgecolor': 'black', 'linewidth': 1.5},
        textprops={'color': 'black', 'weight': 'bold', 'fontsize': 14}
    )
    ax.set_title("\U0001F4CA Resume vs Job Description Match", fontsize=16, weight='bold', color="#2e86de")
    ax.axis('equal')

    mplcursors.cursor(wedges, hover=True).connect(
        "add", lambda sel: sel.annotation.set_text(f'{match_labels[sel.index]}: {match_data[sel.index]}%')
    )

    st.pyplot(fig)

def create_resume_docx(data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.add_heading(data['name'], level=1)
    doc.add_paragraph(f"Email: {data['email']} | Phone: {data['phone']} | LinkedIn: {data['linkedin']}")

    doc.add_heading("Career Objective", level=2)
    doc.add_paragraph(data['objective'])

    doc.add_heading("Education", level=2)
    doc.add_paragraph(data['education'])

    doc.add_heading("Skills", level=2)
    doc.add_paragraph(data['skills'])

    doc.add_heading("Experience", level=2)
    doc.add_paragraph(data['experience'])

    doc.add_heading("Projects", level=2)
    doc.add_paragraph(data['projects'])

    doc.add_heading("Certifications", level=2)
    doc.add_paragraph(data['certifications'])

    doc.add_heading("Languages Known", level=2)
    doc.add_paragraph(data['languages'])

    doc.add_heading("Hobbies", level=2)
    doc.add_paragraph(data['hobbies'])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---- PROMPTS ----
input_prompt1 = """
You are an experienced Technical Human Resource Manager. Your task is to review the provided resume against the job description. 
Please share your professional evaluation on whether the candidate's profile aligns with the role. 
Highlight the strengths and weaknesses of the applicant in relation to the specified job requirements.
"""

input_prompt3 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality. 
Your task is to evaluate the resume against the provided job description. Give me the percentage of match if the resume matches
the job description. First the output should come as percentage and then keywords missing and last final thoughts.
"""

# ---- STREAMLIT CONFIG ----
st.set_page_config(page_title="ResumeCart", layout="centered")

# ---- HEADER ----
st.markdown("<h1 style='text-align: center; color: Red ;'>🛒 ResumeCart</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; font-size: 20px;'>Smart Resume Matching System</p>", unsafe_allow_html=True)

# ---- NAVIGATION ----
menu = st.sidebar.selectbox("Navigation", ["Check Your ATS", "Make Your Resume"])

if menu == "Check Your ATS":
    st.markdown("## 👋 Welcome")

    if "show_single" not in st.session_state:
        st.session_state.show_single = False
    if "show_multiple" not in st.session_state:
        st.session_state.show_multiple = False
    if "result_data" not in st.session_state:
        st.session_state.result_data = None

    st.markdown("### 📋 Provide the Job Description")
    input_text = st.text_area("Paste the job description here:", key="input")

    MAX_CHARS = 8000
    if len(input_text) > MAX_CHARS:
        st.warning(f"⚠️ Job description too long ({len(input_text)} characters). Trimming to {MAX_CHARS}.")
        input_text = input_text[:MAX_CHARS]

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Check for Single Resume"):
            if not input_text.strip():
                st.warning("⚠️ Please enter the job description before proceeding.")
            else:
                st.session_state.show_single = True
                st.session_state.show_multiple = False

    with col2:
        if st.button("Check for Multiple Resumes"):
            if not input_text.strip():
                st.warning("⚠️ Please enter the job description before proceeding.")
            else:
                st.session_state.show_single = False
                st.session_state.show_multiple = True

    if st.session_state.show_single:
        st.markdown("### 📂 Upload Resume")
        uploaded_file = st.file_uploader("Upload your resume (PDF only)", type=["pdf"])

        if uploaded_file is not None:
            st.success("✅ Resume uploaded successfully.")
            col1, col2 = st.columns(2)
            with col1:
                submit1 = st.button("🧠 Tell About The Resume")
            with col2:
                submit3 = st.button("📊 Percentage Match")

            if submit1:
                pdf_content = input_pdf_setup(uploaded_file)
                response = get_gemini_response(input_prompt1, pdf_content, input_text)
                st.markdown("### 📝 Evaluation Result")
                st.write(response)

            elif submit3:
                pdf_content = input_pdf_setup(uploaded_file)
                response = get_gemini_response(input_prompt3, pdf_content, input_text)

                match_percent = extract_percentage(response)
                missing_keywords = extract_missing_keywords(response)
                final_thoughts = response.split("Final thoughts")[-1].strip() if "Final thoughts" in response else "N/A"

                st.markdown("### 📈 Match Analysis Result")
                st.markdown(f"**✅ Match Score: {match_percent}%**")
                show_percentage_chart(match_percent)
                st.markdown("### ❌ Missing Keywords")
                st.write(missing_keywords)
                st.markdown("### 💬 Final Thoughts")
                st.write(final_thoughts)

    if st.session_state.show_multiple:
        st.markdown("### 📂 Multi-Resume Evaluator (Batch Processing)")
        multi_files = st.file_uploader("Upload multiple resumes (PDFs)", type=["pdf"], accept_multiple_files=True)

        if st.button("📊 Evaluate All Resumes") and multi_files:
            result_data = []
            for file in multi_files:
                try:
                    pdf_content = input_pdf_setup(file)
                    response = get_gemini_response(input_prompt3, pdf_content, input_text)

                    match_percent = extract_percentage(response)
                    missing_keywords = extract_missing_keywords(response)
                    final_thoughts = response.split("Final thoughts")[-1].strip() if "Final thoughts" in response else "N/A"

                    result_data.append({
                        "Resume Name": file.name,
                        "Match Score (%)": match_percent,
                        "Missing Keywords": missing_keywords,
                        "Final Thoughts": final_thoughts
                    })
                except Exception as e:
                    result_data.append({
                        "Resume Name": file.name,
                        "Match Score (%)": "Error",
                        "Missing Keywords": str(e),
                        "Final Thoughts": "Processing failed"
                    })
            st.session_state.result_data = result_data

        if st.session_state.result_data:
            df_result = pd.DataFrame(st.session_state.result_data)
            df_result["Match Score (%)"] = pd.to_numeric(df_result["Match Score (%)"], errors='coerce')
            df_result_filtered = df_result.dropna(subset=["Match Score (%)"])
            df_result_filtered["Match Score (%)"] = df_result_filtered["Match Score (%)"].astype(int)

            st.markdown("### 🎯 Filter Resumes by Minimum Match Score")
            min_score = st.selectbox("Select minimum score to filter resumes", options=[50,55,60,65,70,75,80,85,90,95], index=1)
            shortlisted_df = df_result_filtered[df_result_filtered["Match Score (%)"] >= min_score]

            st.markdown("### ✅ Shortlisted Resumes")
            st.dataframe(shortlisted_df.reset_index(drop=True), use_container_width=True)

            with st.expander("🔍 See All Results"):
                st.dataframe(df_result.reset_index(drop=True), use_container_width=True)

elif menu == "Make Your Resume":
    st.markdown("## 📝 Resume Builder")
    st.markdown("Fill in the details below to generate a professional resume.")

    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone Number")
    linkedin = st.text_input("LinkedIn URL")
    objective = st.text_area("Career Objective")
    education = st.text_area("Education")
    skills = st.text_area("Skills")
    experience = st.text_area("Experience")
    projects = st.text_area("Projects")
    certifications = st.text_area("Certifications")
    languages = st.text_area("Languages Known")
    hobbies = st.text_area("Hobbies")

    if st.button("Generate Resume"):
        resume_data = {
            "name": name,
            "email": email,
            "phone": phone,
            "linkedin": linkedin,
            "objective": objective,
            "education": education,
            "skills": skills,
            "experience": experience,  
            "projects": projects,
            "certifications": certifications,
            "languages": languages,
            "hobbies": hobbies
        }
        docx_file = create_resume_docx(resume_data)
        st.success("✅ Resume generated successfully!")
        st.download_button(
            label="📄 Download Resume as DOCX",
            data=docx_file,
            file_name="Generated_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
